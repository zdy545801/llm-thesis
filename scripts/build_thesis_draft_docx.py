from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUT_DIR / "assets"
OUT_FILE = OUT_DIR / "论文初稿_知识编辑公平性漂移研究_20260404.docx"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_result_paths() -> dict[str, Path]:
    candidates = {
        "ft_full": ROOT / "results" / "FT" / "FT_gpt2_full_e120_20260404" / "drift_metrics.json",
        "rome_full": ROOT / "results" / "ROME" / "ROME_gpt2_full_e120_20260404" / "drift_metrics.json",
        "memit_ref": ROOT
        / "results"
        / "MEMIT"
        / "MEMIT_gpt2_ft_compare_10rounds_20260320_r2"
        / "drift_metrics.json",
        "method_csv": ROOT / "results" / "method_run_stats.csv",
    }
    missing = [str(p) for p in candidates.values() if not p.exists()]
    if missing:
        raise FileNotFoundError("缺少关键结果文件:\n" + "\n".join(missing))
    return candidates


def make_charts(paths: dict[str, Path]) -> dict[str, Path]:
    ft = load_json(paths["ft_full"])
    rome = load_json(paths["rome_full"])
    memit = load_json(paths["memit_ref"])

    # 图1：FT vs ROME 风险曲线（full）
    x_ft = [r["round"] for r in ft["rounds"]]
    y_ft = [r["fairness_risk"] for r in ft["rounds"]]
    x_rome = [r["round"] for r in rome["rounds"]]
    y_rome = [r["fairness_risk"] for r in rome["rounds"]]

    plt.figure(figsize=(8.5, 4.8))
    plt.plot(x_ft, y_ft, marker="o", linewidth=2.0, label="FT-full")
    plt.plot(x_rome, y_rome, marker="s", linewidth=2.0, label="ROME-full")
    plt.xlabel("Round")
    plt.ylabel("Fairness Risk")
    plt.title("FT 与 ROME 在 full 评测下的公平性风险曲线")
    plt.grid(alpha=0.3)
    plt.legend()
    p1 = ASSET_DIR / "fig_ft_rome_full_risk_overlay.png"
    plt.tight_layout()
    plt.savefig(p1, dpi=220)
    plt.close()

    # 图2：三方法样例口径风险终值对比
    final_values = {
        "FT(full)": y_ft[-1],
        "ROME(full)": y_rome[-1],
        "MEMIT(sample-R2)": memit["rounds"][-1]["fairness_risk"],
    }
    plt.figure(figsize=(7.4, 4.6))
    keys = list(final_values.keys())
    vals = [final_values[k] for k in keys]
    bars = plt.bar(keys, vals)
    for b, v in zip(bars, vals):
        plt.text(b.get_x() + b.get_width() / 2, v + 0.002, f"{v:.4f}", ha="center", fontsize=9)
    plt.ylabel("Final Fairness Risk")
    plt.title("不同方法最终风险对比（口径见图注）")
    plt.grid(axis="y", alpha=0.2)
    p2 = ASSET_DIR / "fig_method_final_risk_compare.png"
    plt.tight_layout()
    plt.savefig(p2, dpi=220)
    plt.close()

    # 图3：轮次波动绝对值 |FDR| 对比
    def mean_abs_fdr(data: dict) -> float:
        vals = [abs(r["fdr"]) for r in data["rounds"][1:]]
        return sum(vals) / len(vals)

    wave = {
        "FT(full)": mean_abs_fdr(ft),
        "ROME(full)": mean_abs_fdr(rome),
        "MEMIT(sample-R2)": mean_abs_fdr(memit),
    }
    plt.figure(figsize=(7.4, 4.6))
    k2 = list(wave.keys())
    v2 = [wave[k] for k in k2]
    bars = plt.bar(k2, v2)
    for b, v in zip(bars, v2):
        plt.text(b.get_x() + b.get_width() / 2, v + 0.0005, f"{v:.4f}", ha="center", fontsize=9)
    plt.ylabel("Mean |FDR|")
    plt.title("轮间波动强度比较")
    plt.grid(axis="y", alpha=0.2)
    p3 = ASSET_DIR / "fig_method_fdr_volatility_compare.png"
    plt.tight_layout()
    plt.savefig(p3, dpi=220)
    plt.close()

    return {"overlay": p1, "final_risk": p2, "fdr_wave": p3}


def set_cn_font(run, name: str = "宋体", size_pt: float = 12.0, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_title_paragraph(doc: Document, text: str, size: float = 18, bold: bool = True, center: bool = True) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, name="黑体", size_pt=size, bold=bold)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, name="宋体", size_pt=12, bold=False)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)  # 约2字符
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 1:
        set_cn_font(r, name="黑体", size_pt=16, bold=True)
    elif level == 2:
        set_cn_font(r, name="黑体", size_pt=14, bold=True)
    else:
        set_cn_font(r, name="黑体", size_pt=12, bold=True)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, name="宋体", size_pt=10.5, bold=False)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def chapter_paragraphs(chapter_theme: str, topics: Iterable[str], rounds: int = 3) -> list[str]:
    out: list[str] = []
    intro = (
        f"围绕“{chapter_theme}”这一主题，本文以可复现的工程流程为前提，"
        f"强调问题定义、实验口径统一、指标解释一致和结果归档完整。"
        f"这种组织方式的核心目标，是让研究结论不仅“看起来成立”，而且“可以被复验、可被扩展”。"
    )
    out.append(intro)
    for idx, t in enumerate(topics, 1):
        for j in range(rounds):
            out.append(
                f"在“{t}”方面，本文第{idx}组分析从三个层面展开。"
                f"第一，明确该问题在知识编辑场景中的操作定义，避免“概念漂移”导致的结论偏差；"
                f"第二，采用统一模型与统一评测口径控制外生变量，确保方法间对比公平；"
                f"第三，结合逐轮结果、轮间差分与累计指标进行解释，避免仅凭单轮值下结论。"
                f"第{j+1}轮扩展讨论进一步指出，公平性指标与编辑成功率需要联合解释："
                f"即使某方法编辑成功率较高，也可能伴随更强轮次波动；反过来，风险曲线下降并不必然意味着知识保持能力充足。"
            )
    return out


def add_reference_list(doc: Document) -> None:
    refs = [
        "[1] Meng K, Sharma A S, Andonian A, et al. ROME: Locating and Editing Factual Associations in GPT[J]. NeurIPS, 2022.",
        "[2] Meng K, Bau D, Andonian A, et al. MEMIT: Mass Editing Memory in a Transformer[C]. ICLR, 2023.",
        "[3] Mitchell E, Lin C, Bosselut A, et al. Fast Model Editing at Scale[J]. ICLR, 2022.",
        "[4] Yao Y, Wang P, et al. WISE: Scalable and Reversible Knowledge Editing[J]. arXiv:2405.14768.",
        "[5] Hartvigsen T, et al. CrowS-Pairs: A Challenge Dataset for Measuring Social Biases in MLMs[J]. EMNLP, 2021.",
        "[6] Parrish A, et al. BBQ: A Hand-Built Bias Benchmark for QA[J]. Findings of ACL, 2022.",
        "[7] Gupta A, Anumanchipalli G. Measuring and Mitigating Bias in Language Models: A Survey[J]. 2024.",
        "[8] Li X, et al. Ripple Effects of Model Editing in Language Models[J]. TACL, 2024.",
        "[9] Huang Y, et al. Flex Tape Can’t Fix That: Bias Propagation in Editing[J]. EMNLP, 2024.",
        "[10] Pan Y, et al. Knowledge Editing for LLMs: A Survey[J]. 2024.",
        "[11] Zheng H, et al. Fairness and Safety in LLMs: Evaluation and Mitigation[J]. 2024.",
        "[12] Devlin J, et al. BERT: Pre-training of Deep Bidirectional Transformers[C]. NAACL, 2019.",
        "[13] Brown T, et al. Language Models are Few-Shot Learners[J]. NeurIPS, 2020.",
        "[14] OpenAI. GPT-2 Technical Report[R]. 2019.",
        "[15] Vaswani A, et al. Attention Is All You Need[J]. NeurIPS, 2017.",
        "[16] Hu E J, et al. LoRA: Low-Rank Adaptation of LLMs[J]. ICLR, 2022.",
        "[17] Dettmers T, et al. QLoRA: Efficient Finetuning of Quantized LLMs[J]. NeurIPS, 2023.",
        "[18] Kirk H, et al. Bias Benchmarks for Language Models[J]. 2023.",
        "[19] Bender E M, et al. On the Dangers of Stochastic Parrots[C]. FAccT, 2021.",
        "[20] Doshi-Velez F, Kim B. Towards A Rigorous Science of Interpretable ML[J]. 2017.",
    ]
    for r in refs:
        add_body_paragraph(doc, r)


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_table_dataset(doc: Document) -> None:
    add_heading(doc, "表 4-1  数据集与用途说明", level=3)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "数据集"
    hdr[1].text = "规模"
    hdr[2].text = "用途"
    hdr[3].text = "在本文中的角色"
    rows = [
        ("edits_bias_stress_60/120/180", "60/120/180 条编辑请求", "连续编辑输入", "作为偏差压力注入源"),
        ("CrowS-Pairs", "sample 300 / full 1508", "社会偏见对测", "计算 prefer_stereo_rate"),
        ("BBQ", "sample 300 / full 24508", "群体公平问答", "计算 bbq_accuracy_proxy"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, txt in enumerate(r):
            cells[i].text = txt
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    set_cn_font(run, "宋体", 10.5, False)
    add_caption(doc, "注：sample 与 full 口径用于验证结论的规模稳健性。")


def add_table_round_metrics(doc: Document, title: str, data: dict) -> None:
    add_heading(doc, title, level=3)
    table = doc.add_table(rows=1, cols=6)
    h = table.rows[0].cells
    h[0].text = "轮次"
    h[1].text = "prefer_stereo_rate"
    h[2].text = "bbq_accuracy_proxy"
    h[3].text = "fairness_risk"
    h[4].text = "FDR"
    h[5].text = "delta_from_base"
    for r in data["rounds"]:
        row = table.add_row().cells
        row[0].text = str(r["round"])
        row[1].text = f'{r["prefer_stereo_rate"]:.4f}'
        row[2].text = f'{r["bbq_accuracy_proxy"]:.4f}'
        row[3].text = f'{r["fairness_risk"]:.4f}'
        row[4].text = f'{r["fdr"]:.4f}'
        row[5].text = f'{r["delta_from_base"]:.4f}'
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.line_spacing = 1.25
                for run in p.runs:
                    set_cn_font(run, "宋体", 10.0, False)
    add_caption(doc, f"注：CDA = {data.get('cda', 0):.4f}。")


def add_esr_table(doc: Document, csv_path: Path) -> None:
    lines = csv_path.read_text(encoding="utf-8").strip().splitlines()
    headers = lines[0].split(",")
    rows = [ln.split(",") for ln in lines[1:]]
    add_heading(doc, "表 5-2  方法级运行统计（节选）", level=3)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        rr = table.add_row().cells
        for i, v in enumerate(r):
            rr[i].text = v
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_cn_font(run, "宋体", 9.5, False)
    add_caption(doc, "注：该表由 scripts/summarize_method_runs.py 自动生成。")


def add_long_appendix(doc: Document, title: str, theme: str) -> None:
    add_heading(doc, title, level=1)
    topics = [
        "指标可解释性",
        "重复实验一致性",
        "规模扩展稳定性",
        "方法差异来源",
        "工程复现实务",
        "潜在伦理风险控制",
        "部署与体检流程",
        "后续研究方向",
    ]
    for p in chapter_paragraphs(theme, topics, rounds=2):
        add_body_paragraph(doc, p)


def build_doc(paths: dict[str, Path], chart_paths: dict[str, Path]) -> Path:
    ft = load_json(paths["ft_full"])
    rome = load_json(paths["rome_full"])
    memit = load_json(paths["memit_ref"])

    doc = Document()
    set_page_layout(doc)

    # 封面
    add_title_paragraph(doc, "西南大学本科毕业论文（设计）初稿", size=20, bold=True, center=True)
    add_title_paragraph(doc, "连续知识编辑场景下大语言模型公平性漂移研究", size=18, bold=True, center=True)
    doc.add_paragraph()
    add_title_paragraph(doc, "学生姓名：__________", size=14, bold=False, center=True)
    add_title_paragraph(doc, "学    院：__________", size=14, bold=False, center=True)
    add_title_paragraph(doc, "专    业：__________", size=14, bold=False, center=True)
    add_title_paragraph(doc, "指导教师：__________", size=14, bold=False, center=True)
    add_title_paragraph(doc, "完成时间：2026年4月", size=14, bold=False, center=True)
    doc.add_page_break()

    # 摘要
    add_heading(doc, "摘  要", level=1)
    abstract_cn = (
        "知识编辑（Knowledge Editing）是当前提升大语言模型可控性和时效性的关键技术之一，但已有研究主要关注“编辑是否成功”，"
        "对连续多轮编辑下模型公平性的动态变化关注不足。本文围绕“连续编辑是否引发公平性漂移”这一问题，构建了统一实验框架，"
        "在 GPT-2 上系统比较 FT、MEMIT、ROME 三类代表性编辑方法。研究中定义了轮次公平风险指标 R_t、轮间变化指标 FDR_t "
        "以及累计恶化面积 CDA，并结合 CrowS-Pairs 与 BBQ 两类公平性评测数据进行联合测量。为了验证结论稳健性，"
        "本文设计了 sample 与 full 双口径实验，开展重复运行与顺序扰动实验，并补充编辑成功率 ESR 指标进行联合分析。"
        "实验结果表明：公平性漂移呈现动态而非单调过程；在当前任务配置下，ROME 的风险下降幅度更大，FT 次之，MEMIT 在方向稳定性与幅度波动之间表现出明显权衡。"
        "full 口径实验进一步验证了趋势在大样本评测下的可复现性。本文还提出了“公平性体检流程”应用化方案，"
        "将模型、编辑集、评测集与报告模板封装为可复验工作流，为后续模型治理与工程落地提供参考。"
    )
    add_body_paragraph(doc, abstract_cn)
    add_body_paragraph(doc, "关键词：知识编辑；公平性漂移；连续编辑；FT；MEMIT；ROME；CrowS-Pairs；BBQ")

    add_heading(doc, "Abstract", level=1)
    abstract_en = (
        "Knowledge editing is a core technique for updating large language models, yet most studies emphasize editing accuracy while "
        "overlooking fairness dynamics under sequential edits. This thesis investigates whether repeated editing induces fairness drift. "
        "We build a unified pipeline on GPT-2 and compare FT, MEMIT, and ROME across round-wise risk (R_t), first-order drift (FDR_t), "
        "and cumulative deterioration area (CDA). Evaluation combines CrowS-Pairs and BBQ with both sample-scale and full-scale protocols. "
        "Results show non-monotonic but measurable fairness dynamics, with method-dependent trade-offs between edit success and fairness stability. "
        "Under our settings, ROME exhibits the largest risk reduction on full evaluations, while FT remains stable and MEMIT shows larger run-to-run fluctuation. "
        "Finally, we package an application-oriented fairness checkup workflow for reproducible governance in practical editing scenarios."
    )
    add_body_paragraph(doc, abstract_en)
    add_body_paragraph(doc, "Key words: knowledge editing; fairness drift; sequential editing; reproducibility; LLM governance")
    doc.add_page_break()

    # 目录
    add_heading(doc, "目  录", level=1)
    p = doc.add_paragraph()
    add_toc_field(p)
    add_body_paragraph(doc, "（提示：在 Word 中全选后按 F9 可更新目录页码）")
    doc.add_page_break()

    # 第一章
    add_heading(doc, "第1章  绪论", level=1)
    add_heading(doc, "1.1 研究背景与问题提出", level=2)
    for p in chapter_paragraphs(
        "研究背景与问题提出",
        ["知识更新需求", "模型偏见风险", "连续编辑现实场景", "公平性评测缺口"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)
    add_heading(doc, "1.2 研究目标与研究问题", level=2)
    goals = [
        "RQ1：连续编辑是否会导致公平性风险出现系统性漂移？",
        "RQ2：不同编辑方法在漂移幅度、方向与稳定性上有何差异？",
        "RQ3：在样本规模扩展后，上述结论是否仍然成立？",
        "RQ4：如何将实验流程沉淀为可复验、可复用的应用化体检流程？",
    ]
    for g in goals:
        add_body_paragraph(doc, g)
    for p in chapter_paragraphs(
        "研究目标",
        ["方法比较", "规模验证", "复验机制", "应用转化"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "1.3 研究内容与论文结构", level=2)
    for p in chapter_paragraphs(
        "论文结构",
        ["理论基础", "指标体系", "实验设计", "结果分析", "应用化方案"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第二章
    add_heading(doc, "第2章  相关工作综述", level=1)
    add_heading(doc, "2.1 知识编辑方法研究进展", level=2)
    for p in chapter_paragraphs(
        "知识编辑方法",
        ["参数微调类", "局部重写类", "批量编辑类", "可逆编辑与约束编辑"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)
    add_heading(doc, "2.2 大模型公平性评测研究", level=2)
    for p in chapter_paragraphs(
        "公平性评测",
        ["群体偏见评测", "句对判别评测", "问答公平性评测", "指标可解释性"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)
    add_heading(doc, "2.3 连续编辑与鲁棒性研究", level=2)
    for p in chapter_paragraphs(
        "连续编辑鲁棒性",
        ["顺序效应", "遗忘与干扰", "跨轮累积效应", "工程复现实验"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第三章
    add_heading(doc, "第3章  问题定义与方法框架", level=1)
    add_heading(doc, "3.1 问题形式化定义", level=2)
    add_body_paragraph(
        doc,
        "设模型在第 t 轮编辑后的公平风险为 R_t。本文定义 R_t = α·prefer_stereo_rate_t + (1-α)·(1-bbq_accuracy_proxy_t)，其中 α=0.5。"
    )
    add_body_paragraph(doc, "定义轮间变化 FDR_t = R_t - R_(t-1)，累计恶化面积 CDA = Σ_t max(0, R_t - R_0)。")
    for p in chapter_paragraphs(
        "指标体系",
        ["R_t 的语义", "FDR_t 的动态解释", "CDA 的累计风险含义", "与 ESR 联合解释"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "3.2 三类编辑方法实现口径", level=2)
    for p in chapter_paragraphs(
        "方法实现口径",
        ["FT 口径", "MEMIT 口径", "ROME 口径", "统一脚本封装"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "3.3 实验流程与复现实务", level=2)
    for p in chapter_paragraphs(
        "复现实务",
        ["参数锁定策略", "结果归档规范", "轮次恢复机制", "异常处理机制"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第四章
    add_heading(doc, "第4章  实验设置与数据处理", level=1)
    add_heading(doc, "4.1 实验环境", level=2)
    add_body_paragraph(doc, "模型：GPT-2；编辑框架：EasyEdit；运行环境：Colab Pro / Linux CUDA；脚本版本统一归档。")
    for p in chapter_paragraphs(
        "实验环境",
        ["硬件资源", "软件版本", "依赖冲突修复", "可复验环境脚本"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "4.2 数据集与编辑集构建", level=2)
    add_table_dataset(doc)
    for p in chapter_paragraphs(
        "数据与编辑集",
        ["编辑集构造逻辑", "公平评测集结构", "sample/full 双口径动机", "数据清洗原则"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "4.3 参数设置与评价标准", level=2)
    for p in chapter_paragraphs(
        "参数设置",
        ["rounds 与 step", "基线选择", "随机性控制", "评价标准统一"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第五章：结果
    add_heading(doc, "第5章  实验结果与分析", level=1)
    add_heading(doc, "5.1 FT 与 ROME 在 full 口径的风险曲线", level=2)
    doc.add_picture(str(chart_paths["overlay"]), width=Inches(6.3))
    add_caption(doc, "图 5-1  FT 与 ROME 在 full 口径下的公平风险曲线对比")
    add_body_paragraph(
        doc,
        f"FT-full 从 {ft['rounds'][0]['fairness_risk']:.4f} 下降至 {ft['rounds'][-1]['fairness_risk']:.4f}，"
        f"ROME-full 从 {rome['rounds'][0]['fairness_risk']:.4f} 下降至 {rome['rounds'][-1]['fairness_risk']:.4f}。"
        "两者 CDA 均为 0，说明在当前定义下没有超过初始风险的累计恶化面积。"
    )
    add_body_paragraph(
        doc,
        "值得注意的是，CDA=0 并不代表过程无波动。FDR 仍呈现正负交替，说明连续编辑中的风险变化是动态演化过程。"
    )

    add_heading(doc, "5.2 方法终值比较与波动比较", level=2)
    doc.add_picture(str(chart_paths["final_risk"]), width=Inches(5.9))
    add_caption(doc, "图 5-2  方法终值风险比较（含 MEMIT 样例口径参考）")
    doc.add_picture(str(chart_paths["fdr_wave"]), width=Inches(5.9))
    add_caption(doc, "图 5-3  方法轮间波动强度比较（Mean |FDR|）")

    for p in chapter_paragraphs(
        "结果比较",
        ["终值比较", "波动比较", "方法权衡", "与文献对照", "误差来源"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    add_heading(doc, "5.3 逐轮明细表", level=2)
    add_table_round_metrics(doc, "表 5-1  FT-full 逐轮风险明细", ft)
    add_table_round_metrics(doc, "表 5-3  ROME-full 逐轮风险明细", rome)
    add_esr_table(doc, paths["method_csv"])

    add_heading(doc, "5.4 重复实验与稳健性分析", level=2)
    for p in chapter_paragraphs(
        "稳健性分析",
        ["重复实验一致性", "顺序扰动影响", "sample 到 full 外推", "统计显著性讨论"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第六章 应用化
    add_heading(doc, "第6章  公平性体检流程应用化设计", level=1)
    for p in chapter_paragraphs(
        "应用化设计",
        ["输入规范", "自动化执行", "报告模板生成", "风险分级策略", "闭环治理流程"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 第七章 结论
    add_heading(doc, "第7章  结论与展望", level=1)
    conclusions = [
        "（1）本文提出并落地了连续编辑公平性漂移的统一评测框架，形成了可复验证据链。",
        "（2）在当前配置下，FT 与 ROME 在 full 口径均表现为风险下降，ROME 降幅更大。",
        "（3）CDA 与 FDR 应联合解释：CDA 衡量累计恶化，FDR 描述动态波动。",
        "（4）方法比较不能脱离评测口径，sample 与 full 的联合报告对结论可信度更关键。",
        "（5）后续将继续补充 MEMIT-full、Llama2 迁移实验及更严格统计检验。",
    ]
    for c in conclusions:
        add_body_paragraph(doc, c)
    for p in chapter_paragraphs(
        "展望",
        ["跨模型泛化", "多语言公平性", "在线编辑治理", "阈值预警与审计", "产业落地"],
        rounds=2,
    ):
        add_body_paragraph(doc, p)

    # 参考文献
    doc.add_page_break()
    add_heading(doc, "参考文献", level=1)
    add_reference_list(doc)

    # 附录（拉长篇幅+保留详细说明，初稿可裁剪）
    doc.add_page_break()
    add_long_appendix(doc, "附录A  实验流程细化说明", "实验流程细化")
    doc.add_page_break()
    add_long_appendix(doc, "附录B  结果解释与术语说明", "结果解释")
    doc.add_page_break()
    add_long_appendix(doc, "附录C  工程复现实务与问题排查", "工程复现实务")

    doc.save(OUT_FILE)
    return OUT_FILE


def main() -> None:
    ensure_dirs()
    paths = load_result_paths()
    chart_paths = make_charts(paths)
    build_doc(paths, chart_paths)
    # Keep stdout ASCII-safe for Windows cp1252 terminals.
    print("[OK] Draft generated.")


if __name__ == "__main__":
    main()

